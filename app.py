from flask import Flask, render_template, request, send_file, jsonify, abort, redirect, url_for
from PyPDF2 import PdfReader, PdfWriter
from pdf2image import convert_from_path
from docx import Document
import os
import tempfile
from werkzeug.utils import secure_filename
from pdf2docx import Converter
import mimetypes
from docx2pdf import convert
import shutil
import signal
import threading
import time

# Add COM initialization for Windows
try:
    import pythoncom
    COM_INITIALIZED = False
except ImportError:
    pythoncom = None
    COM_INITIALIZED = False

app = Flask(__name__, static_folder='static', static_url_path='/static')

# Configuration
UPLOAD_FOLDER = 'uploads'
OUTPUT_FOLDER = 'output'
ALLOWED_EXTENSIONS = {'pdf', 'docx', 'jpg', 'jpeg', 'png'}

# Set Poppler path
POPPLER_PATH = os.path.join(os.getcwd(), 'poppler', 'poppler-24.02.0', 'Library', 'bin')

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['OUTPUT_FOLDER'] = OUTPUT_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 100 * 1024 * 1024  # 100MB max file size (increased from 16MB)

# Increase timeout for large file processing
app.config['SEND_FILE_MAX_AGE_DEFAULT'] = 0
app.config['PERMANENT_SESSION_LIFETIME'] = 1800  # 30 minutes

# Ensure folders exist
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

# Add MIME type mappings
mimetypes.add_type('application/vnd.openxmlformats-officedocument.wordprocessingml.document', '.docx')
mimetypes.add_type('image/jpeg', '.jpg')
mimetypes.add_type('application/pdf', '.pdf')

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def timeout_handler(signum, frame):
    """Handle timeout for long-running conversions"""
    raise TimeoutError("Conversion timed out")

def run_with_timeout(func, args, timeout_seconds=300):  # 5 minutes timeout
    """Run a function with a timeout"""
    result = [None]
    exception = [None]
    
    def target():
        try:
            result[0] = func(*args)
        except Exception as e:
            exception[0] = e
    
    thread = threading.Thread(target=target)
    thread.daemon = True
    thread.start()
    thread.join(timeout_seconds)
    
    if thread.is_alive():
        # Thread is still running, conversion timed out
        raise TimeoutError(f"Conversion timed out after {timeout_seconds} seconds")
    
    if exception[0]:
        raise exception[0]
    
    return result[0]

def convert_word_to_pdf(input_file, output_file):
    try:
        # Ensure input file exists and is readable
        if not os.path.exists(input_file):
            raise Exception(f"Input file not found: {input_file}")
        
        # Check if input file is readable
        if not os.access(input_file, os.R_OK):
            raise Exception(f"Input file not readable: {input_file}")
        
        # Ensure output directory exists
        output_dir = os.path.dirname(output_file)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)
        
        # Get file size to determine timeout and processing strategy
        file_size = os.path.getsize(input_file)
        timeout_seconds = 300  # Default 5 minutes
        
        # Increase timeout for larger files
        if file_size > 50 * 1024 * 1024:  # 50MB
            timeout_seconds = 900  # 15 minutes
        elif file_size > 20 * 1024 * 1024:  # 20MB
            timeout_seconds = 600  # 10 minutes
        elif file_size > 10 * 1024 * 1024:  # 10MB
            timeout_seconds = 450  # 7.5 minutes
        
        print(f"Converting file: {input_file} (Size: {file_size} bytes, Timeout: {timeout_seconds}s)")
        
        def conversion_task():
            try:
                # Method 1: Try docx2pdf first (requires Microsoft Word)
                try:
                    print("Attempting conversion with docx2pdf...")
                    
                    # Initialize COM for Windows if available
                    global COM_INITIALIZED
                    if pythoncom and not COM_INITIALIZED:
                        try:
                            pythoncom.CoInitialize()
                            COM_INITIALIZED = True
                            print("COM initialized successfully")
                        except Exception as com_error:
                            print(f"COM initialization failed: {str(com_error)}")
                    
                    convert(input_file, output_file)
                    
                    # Verify the output file was created and has content
                    if os.path.exists(output_file) and os.path.getsize(output_file) > 0:
                        print(f"Successfully converted {input_file} to {output_file} using docx2pdf")
                        return True
                    else:
                        raise Exception("docx2pdf conversion failed: Output file is empty or not created")
                        
                except Exception as docx2pdf_error:
                    print(f"docx2pdf failed: {str(docx2pdf_error)}")
                    
                    # Method 2: Try alternative method using python-docx and reportlab
                    try:
                        print("Attempting alternative conversion method with reportlab...")
                        from docx import Document
                        from reportlab.lib.pagesizes import letter
                        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
                        from reportlab.lib.styles import getSampleStyleSheet
                        from reportlab.lib.units import inch
                        from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
                        
                        # Read the Word document
                        doc = Document(input_file)
                        
                        # Create PDF
                        pdf_doc = SimpleDocTemplate(output_file, pagesize=letter)
                        styles = getSampleStyleSheet()
                        story = []
                        
                        # Extract text from paragraphs with basic formatting
                        for paragraph in doc.paragraphs:
                            if paragraph.text.strip():
                                # Determine alignment
                                alignment = TA_LEFT
                                if paragraph.alignment == 1:  # Center
                                    alignment = TA_CENTER
                                elif paragraph.alignment == 2:  # Right
                                    alignment = TA_RIGHT
                                
                                # Create paragraph with alignment
                                style = styles['Normal']
                                style.alignment = alignment
                                p = Paragraph(paragraph.text, style)
                                story.append(p)
                                story.append(Spacer(1, 12))
                        
                        # Build PDF
                        if story:
                            pdf_doc.build(story)
                            
                            # Verify the output file was created and has content
                            if os.path.exists(output_file) and os.path.getsize(output_file) > 0:
                                print(f"Successfully converted {input_file} to {output_file} using reportlab")
                                return True
                            else:
                                raise Exception("Reportlab conversion failed: Output file is empty or not created")
                        else:
                            raise Exception("No content found in Word document")
                            
                    except Exception as reportlab_error:
                        print(f"Reportlab conversion failed: {str(reportlab_error)}")
                        
                        # Method 3: Try using fpdf2 (simpler PDF generation)
                        try:
                            print("Attempting conversion with fpdf2...")
                            from docx import Document
                            from fpdf import FPDF
                            
                            # Read the Word document
                            doc = Document(input_file)
                            
                            # Create PDF
                            pdf = FPDF()
                            pdf.add_page()
                            pdf.set_font("Arial", size=12)
                            
                            # Extract text from paragraphs
                            content_found = False
                            for paragraph in doc.paragraphs:
                                if paragraph.text.strip():
                                    # Handle encoding issues by replacing problematic characters
                                    text = paragraph.text.encode('latin-1', 'replace').decode('latin-1')
                                    pdf.multi_cell(0, 10, text)
                                    content_found = True
                            
                            if content_found:
                                pdf.output(output_file)
                                
                                # Verify the output file was created and has content
                                if os.path.exists(output_file) and os.path.getsize(output_file) > 0:
                                    print(f"Successfully converted {input_file} to {output_file} using fpdf2")
                                    return True
                                else:
                                    raise Exception("FPDF2 conversion failed: Output file is empty or not created")
                            else:
                                raise Exception("No content found in Word document")
                                
                        except Exception as fpdf_error:
                            print(f"FPDF2 conversion failed: {str(fpdf_error)}")
                            
                            # Method 4: Try using unoconv if available (LibreOffice)
                            try:
                                print("Attempting conversion with unoconv...")
                                import subprocess
                                
                                # Check if unoconv is available
                                result = subprocess.run(['unoconv', '--version'], capture_output=True, text=True)
                                if result.returncode == 0:
                                    # Use unoconv for conversion
                                    result = subprocess.run(['unoconv', '-f', 'pdf', '-o', output_file, input_file], 
                                                          capture_output=True, text=True, timeout=timeout_seconds)
                                    
                                    if result.returncode == 0 and os.path.exists(output_file) and os.path.getsize(output_file) > 0:
                                        print(f"Successfully converted {input_file} to {output_file} using unoconv")
                                        return True
                                    else:
                                        raise Exception(f"unoconv conversion failed: {result.stderr}")
                                else:
                                    raise Exception("unoconv not available")
                                    
                            except Exception as unoconv_error:
                                print(f"unoconv conversion failed: {str(unoconv_error)}")
                                
                                # If all methods fail, provide a helpful error message
                                error_msg = f"All conversion methods failed. Please ensure you have one of the following:\n"
                                error_msg += "1. Microsoft Word installed (for docx2pdf)\n"
                                error_msg += "2. LibreOffice installed (for unoconv)\n"
                                error_msg += "3. Or try converting a different file format"
                                raise Exception(error_msg)
                    
            except Exception as conversion_error:
                print(f"Conversion error: {str(conversion_error)}")
                # Clean up any partial output
                if os.path.exists(output_file):
                    try:
                        os.remove(output_file)
                    except:
                        pass
                raise conversion_error
        
        # Run conversion with timeout
        result = run_with_timeout(conversion_task, (), timeout_seconds)
        return result
            
    except TimeoutError as te:
        print(f"Conversion timed out: {str(te)}")
        # Clean up any partial output
        if os.path.exists(output_file):
            try:
                os.remove(output_file)
            except:
                pass
        raise Exception(f"Conversion timed out after {timeout_seconds} seconds. Please try with a smaller file or check your internet connection.")
    except Exception as e:
        print(f"Error in convert_word_to_pdf: {str(e)}")
        # Clean up any partial output
        if os.path.exists(output_file):
            try:
                os.remove(output_file)
            except:
                pass
        return False

def convert_pdf_to_word(input_file, output_file):
    try:
        # Ensure input file exists and is readable
        if not os.path.exists(input_file):
            raise Exception(f"Input file not found: {input_file}")
        
        # Check if input file is readable
        if not os.access(input_file, os.R_OK):
            raise Exception(f"Input file not readable: {input_file}")
        
        # Ensure output directory exists
        output_dir = os.path.dirname(output_file)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)
        
        # Get file size to determine timeout and processing strategy
        file_size = os.path.getsize(input_file)
        timeout_seconds = 300  # Default 5 minutes
        
        # Increase timeout for larger files
        if file_size > 50 * 1024 * 1024:  # 50MB
            timeout_seconds = 900  # 15 minutes
        elif file_size > 20 * 1024 * 1024:  # 20MB
            timeout_seconds = 600  # 10 minutes
        elif file_size > 10 * 1024 * 1024:  # 10MB
            timeout_seconds = 450  # 7.5 minutes
        
        print(f"Converting file: {input_file} (Size: {file_size} bytes, Timeout: {timeout_seconds}s)")
        
        def conversion_task():
            # Use pdf2docx for conversion with explicit file handling
            cv = None
            try:
                cv = Converter(input_file)
                cv.convert(output_file)
                
                # Verify the output file exists and has content
                if os.path.exists(output_file) and os.path.getsize(output_file) > 0:
                    print(f"Successfully converted {input_file} to {output_file}")
                    return True
                else:
                    raise Exception("Conversion failed: Output file is empty or not created")
                    
            except Exception as conversion_error:
                print(f"Conversion error: {str(conversion_error)}")
                # Clean up any partial output
                if os.path.exists(output_file):
                    try:
                        os.remove(output_file)
                    except:
                        pass
                raise conversion_error
            finally:
                # Always close the converter
                if cv:
                    try:
                        cv.close()
                    except:
                        pass
        
        # Run conversion with timeout
        result = run_with_timeout(conversion_task, (), timeout_seconds)
        return result
            
    except TimeoutError as te:
        print(f"Conversion timed out: {str(te)}")
        # Clean up any partial output
        if os.path.exists(output_file):
            try:
                os.remove(output_file)
            except:
                pass
        raise Exception(f"Conversion timed out after {timeout_seconds} seconds. Please try with a smaller file or check your internet connection.")
    except Exception as e:
        print(f"Error in PDF to Word conversion: {str(e)}")
        # Clean up any partial output
        if os.path.exists(output_file):
            try:
                os.remove(output_file)
            except:
                pass
        return False

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/about')
def about():
    return render_template('about.html')

@app.route('/select_file')
def select_file():
    conversion_type = request.args.get('type', 'pdf')
    return render_template('select_file.html', type=conversion_type)

@app.route('/convert', methods=['POST'])
def convert_file():
    print("=== Starting file conversion ===")
    
    if 'file' not in request.files:
        print("No file in request")
        return jsonify({'error': 'No file uploaded'}), 400
    
    file = request.files['file']
    conversion_type = request.form.get('type', 'word')
    
    print(f"File: {file.filename}, Type: {conversion_type}")
    
    if file.filename == '':
        print("No filename")
        return jsonify({'error': 'No file selected'}), 400
    
    # Validate file type
    if conversion_type == 'word' and not file.filename.lower().endswith('.pdf'):
        return jsonify({'error': 'Please upload a PDF file for conversion to Word'}), 400
    elif conversion_type == 'pdf' and not file.filename.lower().endswith(('.docx', '.doc')):
        return jsonify({'error': 'Please upload a Word file for conversion to PDF'}), 400
    
    try:
        # Save uploaded file directly to uploads folder
        filename = secure_filename(file.filename)
        input_path = os.path.join(UPLOAD_FOLDER, filename)
        
        print(f"Saving file to: {input_path}")
        file.save(input_path)
        
        # Check if file was saved
        if not os.path.exists(input_path):
            raise Exception("Failed to save uploaded file")
        
        file_size = os.path.getsize(input_path)
        print(f"File saved: {file_size} bytes")
        
        if file_size == 0:
            raise Exception("Uploaded file is empty")
        
        # Check file size limit (100MB)
        if file_size > 100 * 1024 * 1024:
            os.remove(input_path)  # Clean up
            raise Exception("File size exceeds 100MB limit. Please use a smaller file.")
        
        # Determine output filename
        output_filename = os.path.splitext(filename)[0]
        if conversion_type == 'pdf':
            output_filename += '.pdf'
        else:
            output_filename += '.docx'
        
        output_path = os.path.join(UPLOAD_FOLDER, output_filename)
        print(f"Output path: {output_path}")
        
        # Show processing message for large files
        if file_size > 10 * 1024 * 1024:  # 10MB
            print(f"Processing large file ({file_size / (1024*1024):.1f} MB)...")
        
        # Perform conversion with timeout handling
        success = False
        try:
            if conversion_type == 'pdf':
                success = convert_word_to_pdf(input_path, output_path)
            else:
                success = convert_pdf_to_word(input_path, output_path)
        except Exception as conversion_error:
            print(f"Conversion error: {str(conversion_error)}")
            # Clean up input file
            try:
                os.remove(input_path)
            except:
                pass
            raise Exception(f"Conversion failed: {str(conversion_error)}")
        
        print(f"Conversion success: {success}")
        
        if success and os.path.exists(output_path):
            # Get file size
            output_size = os.path.getsize(output_path)
            if output_size < 1024:
                size_str = f"{output_size} B"
            elif output_size < 1024 * 1024:
                size_str = f"{output_size / 1024:.1f} KB"
            else:
                size_str = f"{output_size / (1024 * 1024):.1f} MB"
            
            # Create success message
            if conversion_type == 'pdf':
                success_message = 'Word document converted to PDF successfully!'
            else:
                success_message = 'PDF converted to Word document successfully!'
            
            # Clean up input file
            try:
                os.remove(input_path)
                print(f"Removed input file: {input_path}")
            except:
                pass
            
            # Prepare files info for download page
            files_info = [{'name': output_filename, 'size': size_str}]
            files_param = str(files_info).replace("'", '"')
            
            print(f"Conversion completed: {output_filename}")
            return redirect(url_for('download', success='true', message=success_message, files=files_param))
        else:
            # Clean up input file on failure
            try:
                os.remove(input_path)
            except:
                pass
            raise Exception("Conversion failed. Please try again with a different file.")
            
    except Exception as e:
        print(f"Error in convert_file: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/upload_progress', methods=['POST'])
def upload_progress():
    """Handle large file uploads with progress tracking"""
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file uploaded'}), 400
        
        file = request.files['file']
        conversion_type = request.form.get('type', 'word')
        
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        # Validate file type
        if conversion_type == 'word' and not file.filename.lower().endswith('.pdf'):
            return jsonify({'error': 'Please upload a PDF file for conversion to Word'}), 400
        elif conversion_type == 'pdf' and not file.filename.lower().endswith(('.docx', '.doc')):
            return jsonify({'error': 'Please upload a Word file for conversion to PDF'}), 400
        
        # Save file with progress tracking
        filename = secure_filename(file.filename)
        input_path = os.path.join(UPLOAD_FOLDER, filename)
        
        # Save file in chunks for large files
        chunk_size = 8192  # 8KB chunks
        total_size = 0
        
        with open(input_path, 'wb') as f:
            while True:
                chunk = file.stream.read(chunk_size)
                if not chunk:
                    break
                f.write(chunk)
                total_size += len(chunk)
                
                # Check size limit during upload
                if total_size > 100 * 1024 * 1024:  # 100MB limit
                    f.close()
                    os.remove(input_path)
                    return jsonify({'error': 'File size exceeds 100MB limit'}), 400
        
        if total_size == 0:
            os.remove(input_path)
            return jsonify({'error': 'Uploaded file is empty'}), 400
        
        # Perform conversion
        output_filename = os.path.splitext(filename)[0]
        if conversion_type == 'pdf':
            output_filename += '.pdf'
        else:
            output_filename += '.docx'
        
        output_path = os.path.join(UPLOAD_FOLDER, output_filename)
        
        success = False
        if conversion_type == 'pdf':
            success = convert_word_to_pdf(input_path, output_path)
        else:
            success = convert_pdf_to_word(input_path, output_path)
        
        if success and os.path.exists(output_path):
            # Clean up input file
            try:
                os.remove(input_path)
            except:
                pass
            
            # Get output file size
            output_size = os.path.getsize(output_path)
            if output_size < 1024:
                size_str = f"{output_size} B"
            elif output_size < 1024 * 1024:
                size_str = f"{output_size / 1024:.1f} KB"
            else:
                size_str = f"{output_size / (1024 * 1024):.1f} MB"
            
            files_info = [{'name': output_filename, 'size': size_str}]
            
            return jsonify({
                'success': True,
                'message': f'File converted successfully!',
                'files': files_info
            })
        else:
            # Clean up on failure
            try:
                os.remove(input_path)
            except:
                pass
            return jsonify({'error': 'Conversion failed. Please try again.'}), 500
            
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/split', methods=['POST'])
def split_pdf():
    if 'file' not in request.files:
        return jsonify({'error': 'No file uploaded'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    
    if not allowed_file(file.filename):
        return jsonify({'error': 'Invalid file type'}), 400
    
    filename = secure_filename(file.filename)
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(filepath)
    
    try:
        # Split PDF logic
        pdf = PdfReader(filepath)
        output_files = []
        
        for i, page in enumerate(pdf.pages):
            writer = PdfWriter()
            writer.add_page(page)
            output_filename = f"page_{i+1}.pdf"
            output_path = os.path.join(app.config['UPLOAD_FOLDER'], output_filename)
            
            with open(output_path, 'wb') as out:
                writer.write(out)
            output_files.append(output_filename)
        
        # Clean up input file
        os.remove(filepath)
        
        # Redirect to download page with success message
        files_info = []
        for filename in output_files:
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            if os.path.exists(file_path):
                file_size = os.path.getsize(file_path)
                # Convert to human readable format
                if file_size < 1024:
                    size_str = f"{file_size} B"
                elif file_size < 1024 * 1024:
                    size_str = f"{file_size / 1024:.1f} KB"
                else:
                    size_str = f"{file_size / (1024 * 1024):.1f} MB"
                files_info.append({'name': filename, 'size': size_str})
            else:
                files_info.append({'name': filename, 'size': 'Ready for download'})
        
        files_param = str(files_info).replace("'", '"')
        return redirect(url_for('download', success='true', message='PDF split successfully', files=files_param))
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/to-images', methods=['POST'])
def to_images():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
    
    if not file.filename.lower().endswith('.pdf'):
        return jsonify({'error': 'File must be a PDF'}), 400

    try:
        # Save the uploaded file
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)

        # Convert PDF to images
        images = convert_from_path(filepath, poppler_path=POPPLER_PATH)
        
        # Save images
        image_files = []
        for i, image in enumerate(images):
            image_filename = f"{os.path.splitext(filename)[0]}_page_{i+1}.jpg"
            image_path = os.path.join(app.config['UPLOAD_FOLDER'], image_filename)
            image.save(image_path, 'JPEG')
            image_files.append(image_filename)

        # Clean up the original PDF
        os.remove(filepath)

        # Get file sizes for download page
        files_info = []
        for image_filename in image_files:
            image_path = os.path.join(app.config['UPLOAD_FOLDER'], image_filename)
            if os.path.exists(image_path):
                file_size = os.path.getsize(image_path)
                if file_size < 1024:
                    size_str = f"{file_size} B"
                elif file_size < 1024 * 1024:
                    size_str = f"{file_size / 1024:.1f} KB"
                else:
                    size_str = f"{file_size / (1024 * 1024):.1f} MB"
                files_info.append({'name': image_filename, 'size': size_str})
            else:
                files_info.append({'name': image_filename, 'size': 'Ready for download'})
        
        files_param = str(files_info).replace("'", '"')
        return redirect(url_for('download', success='true', message=f'Successfully converted PDF to {len(image_files)} images', files=files_param))

    except Exception as e:
        return redirect(url_for('download', success='false', message=f'Conversion failed: {str(e)}', files='[]'))

@app.route('/to-pdf', methods=['POST'])
def to_pdf():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
    
    if not file.filename.lower().endswith(('.docx', '.doc')):
        return jsonify({'error': 'File must be a Word document (.docx or .doc)'}), 400

    try:
        # Save the uploaded file
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)

        # Generate output filename
        output_filename = f"{os.path.splitext(filename)[0]}.pdf"
        output_path = os.path.join(app.config['UPLOAD_FOLDER'], output_filename)

        # Convert Word to PDF
        convert(filepath, output_path)

        # Clean up the original Word file
        os.remove(filepath)

        # Get file size for download page
        file_size = os.path.getsize(output_path)
        if file_size < 1024:
            size_str = f"{file_size} B"
        elif file_size < 1024 * 1024:
            size_str = f"{file_size / 1024:.1f} KB"
        else:
            size_str = f"{file_size / (1024 * 1024):.1f} MB"
        
        files_info = [{'name': output_filename, 'size': size_str}]
        files_param = str(files_info).replace("'", '"')
        
        return redirect(url_for('download', success='true', message='Word document converted to PDF successfully!', files=files_param))

    except Exception as e:
        return redirect(url_for('download', success='false', message=f'Conversion failed: {str(e)}', files='[]'))

@app.route('/to-word', methods=['POST'])
def to_word():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
    
    if not file.filename.lower().endswith('.pdf'):
        return jsonify({'error': 'File must be a PDF'}), 400

    try:
        # Save the uploaded file
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)

        # Generate output filename
        output_filename = f"{os.path.splitext(filename)[0]}.docx"
        output_path = os.path.join(app.config['UPLOAD_FOLDER'], output_filename)

        # Convert PDF to Word
        cv = Converter(filepath)
        cv.convert(output_path)
        cv.close()

        # Clean up the original PDF
        os.remove(filepath)

        # Get file size for download page
        file_size = os.path.getsize(output_path)
        if file_size < 1024:
            size_str = f"{file_size} B"
        elif file_size < 1024 * 1024:
            size_str = f"{file_size / 1024:.1f} KB"
        else:
            size_str = f"{file_size / (1024 * 1024):.1f} MB"
        
        files_info = [{'name': output_filename, 'size': size_str}]
        files_param = str(files_info).replace("'", '"')
        
        return redirect(url_for('download', success='true', message='PDF converted to Word document successfully!', files=files_param))

    except Exception as e:
        return redirect(url_for('download', success='false', message=f'Conversion failed: {str(e)}', files='[]'))

@app.route('/download')
def download():
    success = request.args.get('success', 'false') == 'true'
    message = request.args.get('message', '')
    files = request.args.get('files', '[]')
    
    try:
        files = eval(files)  # Convert string representation of list to actual list
    except:
        files = []
    
    return render_template('download.html', success=success, message=message, files=files)

@app.route('/download_file/<filename>')
def download_file(filename):
    """Download a file with improved error handling and streaming for large files"""
    try:
        print(f"=== Download request for: {filename} ===")
        
        # Secure the filename
        safe_filename = secure_filename(filename)
        print(f"Safe filename: {safe_filename}")
        
        # Look for file in uploads folder first, then output folder
        file_path = None
        for folder in [UPLOAD_FOLDER, OUTPUT_FOLDER]:
            test_path = os.path.join(folder, safe_filename)
            if os.path.exists(test_path):
                file_path = test_path
                print(f"File found in {folder}: {file_path}")
                break
        
        if not file_path:
            print(f"File not found in any folder: {safe_filename}")
            return jsonify({'error': f'File "{safe_filename}" not found'}), 404
        
        # Check if file is readable
        if not os.access(file_path, os.R_OK):
            print(f"File not readable: {file_path}")
            return jsonify({'error': f'File "{safe_filename}" is not accessible'}), 403
        
        # Get file size
        try:
            file_size = os.path.getsize(file_path)
            print(f"File size: {file_size} bytes")
            if file_size == 0:
                print("File is empty")
                return jsonify({'error': f'File "{safe_filename}" is empty'}), 400
        except OSError as e:
            print(f"Error getting file size: {e}")
            return jsonify({'error': f'Cannot access file "{safe_filename}"'}), 500
        
        # Determine MIME type
        file_ext = os.path.splitext(safe_filename)[1].lower()
        if file_ext == '.pdf':
            mimetype = 'application/pdf'
        elif file_ext in ['.docx', '.doc']:
            mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        elif file_ext in ['.jpg', '.jpeg']:
            mimetype = 'image/jpeg'
        elif file_ext == '.png':
            mimetype = 'image/png'
        else:
            mimetype = 'application/octet-stream'
        
        print(f"MIME type: {mimetype}")
        
        # For large files (>50MB), use streaming
        if file_size > 50 * 1024 * 1024:  # 50MB
            print("Using streaming for large file download")
            return send_file(
                file_path,
                as_attachment=True,
                download_name=safe_filename,
                mimetype=mimetype,
                conditional=True,  # Enable conditional requests
                etag=True  # Enable ETag for caching
            )
        else:
            # For smaller files, use regular send_file
            response = send_file(
                file_path,
                as_attachment=True,
                download_name=safe_filename,
                mimetype=mimetype
            )
            
            # Add headers to prevent caching issues
            response.headers['Content-Disposition'] = f'attachment; filename="{safe_filename}"'
            response.headers['Cache-Control'] = 'no-cache, no-store, must-revalidate'
            response.headers['Pragma'] = 'no-cache'
            response.headers['Expires'] = '0'
            response.headers['Content-Length'] = str(file_size)
            
            print(f"Download response prepared successfully")
            return response
        
    except Exception as e:
        print(f"Download error: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'Download failed: {str(e)}'}), 500

@app.route('/check_file/<filename>')
def check_file(filename):
    """Check if a file exists and is accessible"""
    try:
        safe_filename = secure_filename(filename)
        
        # Check in both folders
        upload_path = os.path.join(UPLOAD_FOLDER, safe_filename)
        output_path = os.path.join(OUTPUT_FOLDER, safe_filename)
        
        file_info = {
            'exists': False,
            'readable': False,
            'size': 0,
            'path': None,
            'error': None
        }
        
        # Check uploads folder first
        if os.path.exists(upload_path):
            file_info['exists'] = True
            file_info['path'] = upload_path
        elif os.path.exists(output_path):
            file_info['exists'] = True
            file_info['path'] = output_path
        
        if file_info['exists']:
            try:
                # Check if file is readable
                if os.access(file_info['path'], os.R_OK):
                    file_info['readable'] = True
                    file_info['size'] = os.path.getsize(file_info['path'])
                else:
                    file_info['error'] = 'File not readable'
            except Exception as e:
                file_info['error'] = str(e)
        else:
            file_info['error'] = 'File not found'
        
        return jsonify(file_info)
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

def cleanup_old_files():
    """Clean up old files from upload and output folders"""
    try:
        import time
        current_time = time.time()
        max_age = 3600  # 1 hour
        
        for folder in [UPLOAD_FOLDER, OUTPUT_FOLDER]:
            if os.path.exists(folder):
                for filename in os.listdir(folder):
                    file_path = os.path.join(folder, filename)
                    try:
                        # Check file age
                        file_age = current_time - os.path.getmtime(file_path)
                        if file_age > max_age:
                            os.remove(file_path)
                            print(f"Cleaned up old file: {file_path}")
                    except Exception as e:
                        print(f"Error cleaning up {file_path}: {e}")
    except Exception as e:
        print(f"Error in cleanup: {e}")

# Clean up old files on startup
cleanup_old_files()

@app.route('/list_files')
def list_files():
    """List all files in uploads and output folders for debugging"""
    try:
        files_info = {
            'uploads': [],
            'output': []
        }
        
        # List files in uploads folder
        if os.path.exists(UPLOAD_FOLDER):
            for filename in os.listdir(UPLOAD_FOLDER):
                file_path = os.path.join(UPLOAD_FOLDER, filename)
                if os.path.isfile(file_path):
                    size = os.path.getsize(file_path)
                    files_info['uploads'].append({
                        'name': filename,
                        'size': size,
                        'readable': os.access(file_path, os.R_OK)
                    })
        
        # List files in output folder
        if os.path.exists(OUTPUT_FOLDER):
            for filename in os.listdir(OUTPUT_FOLDER):
                file_path = os.path.join(OUTPUT_FOLDER, filename)
                if os.path.isfile(file_path):
                    size = os.path.getsize(file_path)
                    files_info['output'].append({
                        'name': filename,
                        'size': size,
                        'readable': os.access(file_path, os.R_OK)
                    })
        
        return jsonify(files_info)
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/privacy')
def privacy():
    return render_template('privacy.html')

@app.route('/contact', methods=['GET', 'POST'])
def contact():
    success = False
    if request.method == 'POST':
        # You can process/store the message here if needed
        success = True
    return render_template('contact.html', success=success)

@app.route('/terms')
def terms():
    return render_template('terms.html')

@app.route('/conversion_status/<filename>')
def conversion_status(filename):
    """Check the status of a file conversion"""
    try:
        safe_filename = secure_filename(filename)
        
        # Check if input file exists
        input_path = os.path.join(UPLOAD_FOLDER, safe_filename)
        
        # Determine output filename
        base_name = os.path.splitext(safe_filename)[0]
        output_filename = base_name + '.pdf' if safe_filename.lower().endswith(('.docx', '.doc')) else base_name + '.docx'
        output_path = os.path.join(UPLOAD_FOLDER, output_filename)
        
        status = {
            'input_exists': os.path.exists(input_path),
            'output_exists': os.path.exists(output_path),
            'input_size': os.path.getsize(input_path) if os.path.exists(input_path) else 0,
            'output_size': os.path.getsize(output_path) if os.path.exists(output_path) else 0,
            'status': 'processing'
        }
        
        if status['output_exists'] and status['output_size'] > 0:
            status['status'] = 'completed'
        elif not status['input_exists']:
            status['status'] = 'error'
        else:
            status['status'] = 'processing'
        
        return jsonify(status)
        
    except Exception as e:
        return jsonify({'error': str(e), 'status': 'error'}), 500

@app.route('/convert_async', methods=['POST'])
def convert_async():
    """Handle large file conversions asynchronously"""
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file uploaded'}), 400
        
        file = request.files['file']
        conversion_type = request.form.get('type', 'word')
        
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        # Validate file type
        if conversion_type == 'word' and not file.filename.lower().endswith('.pdf'):
            return jsonify({'error': 'Please upload a PDF file for conversion to Word'}), 400
        elif conversion_type == 'pdf' and not file.filename.lower().endswith(('.docx', '.doc')):
            return jsonify({'error': 'Please upload a Word file for conversion to PDF'}), 400
        
        # Save uploaded file
        filename = secure_filename(file.filename)
        input_path = os.path.join(UPLOAD_FOLDER, filename)
        file.save(input_path)
        
        file_size = os.path.getsize(input_path)
        if file_size == 0:
            os.remove(input_path)
            return jsonify({'error': 'Uploaded file is empty'}), 400
        
        if file_size > 100 * 1024 * 1024:
            os.remove(input_path)
            return jsonify({'error': 'File size exceeds 100MB limit'}), 400
        
        # Start conversion in background
        def background_conversion():
            try:
                output_filename = os.path.splitext(filename)[0]
                if conversion_type == 'pdf':
                    output_filename += '.pdf'
                else:
                    output_filename += '.docx'
                
                output_path = os.path.join(UPLOAD_FOLDER, output_filename)
                
                success = False
                if conversion_type == 'pdf':
                    success = convert_word_to_pdf(input_path, output_path)
                else:
                    success = convert_pdf_to_word(input_path, output_path)
                
                if success and os.path.exists(output_path):
                    # Clean up input file
                    try:
                        os.remove(input_path)
                    except:
                        pass
                    print(f"Background conversion completed: {output_filename}")
                else:
                    print(f"Background conversion failed: {filename}")
                    
            except Exception as e:
                print(f"Background conversion error: {str(e)}")
                # Clean up on error
                try:
                    os.remove(input_path)
                except:
                    pass
        
        # Start background thread
        import threading
        thread = threading.Thread(target=background_conversion)
        thread.daemon = True
        thread.start()
        
        return jsonify({
            'success': True,
            'message': 'Conversion started',
            'filename': filename,
            'status_url': f'/conversion_status/{filename}'
        })
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    app.run(debug=True) 